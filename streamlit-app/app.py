import streamlit as st
import requests
import pandas as pd
import plotly.express as px
from datetime import datetime, timedelta
import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches
from PIL import Image

# LOGIN

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if not st.session_state.logged_in:
    st.title("🔐 Sistema de Reportes - Inicio de Sesión")
    usuario = st.text_input("Usuario")
    clave = st.text_input("Contraseña", type="password")
    if st.button("Ingresar"):
        if usuario == "admin" and clave == "pass":
            st.session_state.logged_in = True
            st.rerun()
        else:
            st.error("❌ Usuario o contraseña incorrectos")
    st.stop()


# boton cerrar ses y link a daashboar de grafana

col1, col2, col3 = st.columns([6, 2, 2])  
with col2:
    if st.button("🚪 Cerrar Sesión", key="logout"):
        st.session_state.logged_in = False
        st.rerun()
with col3:
    st.link_button("🌐 Ir a Grafana", 
                   "http://54.234.250.250:3000/d/43a0d842-919e-4934-a8a8-6eb61d472010/windows-exporter-dashboard-20230531-starsl-cn2?orgId=1")


# cobnnfig var como diccionairio

PROMETHEUS_URL = 'http://54.234.250.250:9090'

METRICAS_POR_DEFECTO = [
    "windows_cpu_time_total",
    "windows_memory_physical_free_bytes",
    "windows_logical_disk_free_bytes",
    "windows_net_bytes_received_total",
    "windows_net_bytes_sent_total"
]

ETIQUETAS_METRICAS = {
    "windows_cpu_time_total": "CPU (%)",
    "windows_memory_physical_free_bytes": "RAM Usada (%)",
    "windows_logical_disk_free_bytes": "Disco Usado (%)",
    "windows_net_bytes_received_total": "Red - Recibido (KB/s)",
    "windows_net_bytes_sent_total": "Red - Enviado (KB/s)"
}

reverse_map = {v: k for k, v in ETIQUETAS_METRICAS.items()}


# traer instancia desde Prometheus

@st.cache_data
def obtener_instancias():
    try:
        resp = requests.get(f"{PROMETHEUS_URL}/api/v1/targets", timeout=10)
        data = resp.json()
        return sorted(list(set(
            [t["labels"]["instance"] for t in data["data"]["activeTargets"] if "instance" in t["labels"]]
        )))
    except:
        return ["18.209.1.204:9182"]

#  menú lateral sidebar

st.title("📊 Generar Reporte")
st.sidebar.header("⚙️ Configuración del Reporte")

instancias = obtener_instancias()

opciones_tiempo = {
    "Última hora": timedelta(hours=1),
    "Últimas 6 horas": timedelta(hours=6),
    "Últimas 12 horas": timedelta(hours=12),
    "Últimas 24 horas": timedelta(hours=24),
    "Última semana": timedelta(days=7),
}
rango = st.sidebar.selectbox("Rango de Tiempo", list(opciones_tiempo.keys()))
fin = datetime.now()
inicio = fin - opciones_tiempo[rango]
inicio_ts, fin_ts = int(inicio.timestamp()), int(fin.timestamp())

paso = st.sidebar.selectbox("Intervalo de Muestreo", ["1m", "5m", "15m"], index=1)


# graficos

if "graficos" not in st.session_state:
    st.session_state.graficos = [{
        "metricas": ["windows_cpu_time_total"],
        "instancia": None,
        "tipo": "Línea"
    }]

opciones_metricas = [ETIQUETAS_METRICAS.get(m, m) for m in METRICAS_POR_DEFECTO]

for idx, grafico in enumerate(st.session_state.graficos):
    with st.sidebar.expander(f"📊 Configuración del Gráfico {idx+1}", expanded=True):

        # Instancia por gráfico
        grafico["instancia"] = st.selectbox(
            "Instancia",
            instancias,
            index=instancias.index(grafico["instancia"]) if grafico["instancia"] in instancias else 0,
            key=f"instancia_{idx}"
        )

        # Tipo de gráfico por gráfico
        grafico["tipo"] = st.radio(
            "Tipo de gráfico",
            ["Línea", "Barras"],
            index=0 if grafico["tipo"] == "Línea" else 1,
            key=f"tipo_{idx}"
        )

        # Métricas del gráfico
        for i, m in enumerate(grafico["metricas"]):
            col1, col2 = st.columns([3, 1])
            etiqueta_actual = ETIQUETAS_METRICAS.get(m, m)

            seleccion = col1.selectbox(
                f"➤ Métrica {i+1}",
                opciones_metricas,
                index=opciones_metricas.index(etiqueta_actual) if etiqueta_actual in opciones_metricas else 0,
                key=f"metric_{idx}_{i}"
            )
            grafico["metricas"][i] = reverse_map.get(seleccion, seleccion)

            if col2.button("❌", key=f"remove_{idx}_{i}"):
                grafico["metricas"].pop(i)
                st.rerun()

        st.markdown("---")
        if st.button(f"➕ Agregar métrica a este gráfico", key=f"add_metric_{idx}"):
            grafico["metricas"].append("windows_cpu_time_total")
            st.rerun()

        if st.button(f"🗑️ Eliminar este gráfico", key=f"remove_chart_{idx}"):
            st.session_state.graficos.pop(idx)
            st.rerun()

if st.sidebar.button("➕ Agregar Nuevo Gráfico"):
    st.session_state.graficos.append({
        "metricas": ["windows_cpu_time_total"],
        "instancia": instancias[0] if instancias else None,
        "tipo": "Línea"
    })
    st.rerun()


# consulta a prometheus igual que en grafana

imagenes_graficos = []

if st.sidebar.button("Generar Gráficos"):
    for idx, grafico in enumerate(st.session_state.graficos):
        metricas = grafico["metricas"]
        instancia = grafico["instancia"]
        tipo_grafico = grafico["tipo"]

        dfs = []
        for metrica in metricas:

            if metrica == "windows_cpu_time_total":
                consulta = f'100 - avg(irate(windows_cpu_time_total{{instance="{instancia}", mode="idle"}}[5m])) * 100'
            elif metrica == "windows_memory_physical_free_bytes":
                consulta = f'(1 - (windows_os_physical_memory_free_bytes{{instance="{instancia}"}} / windows_cs_physical_memory_bytes{{instance="{instancia}"}})) * 100'
            elif metrica == "windows_logical_disk_free_bytes":
                consulta = f'(1 - (windows_logical_disk_free_bytes{{instance="{instancia}", volume!~"HarddiskVolume.+"}} / windows_logical_disk_size_bytes{{instance="{instancia}", volume!~"HarddiskVolume.+"}})) * 100'
            elif metrica in ["windows_net_bytes_received_total", "windows_net_bytes_sent_total"]:
                consulta = f'rate({metrica}{{instance="{instancia}"}}[5m]) / 1024'
            else:
                consulta = f'{metrica}{{instance="{instancia}"}}'

            resp = requests.get(f"{PROMETHEUS_URL}/api/v1/query_range", params={
                "query": consulta, "start": inicio_ts, "end": fin_ts, "step": paso
            }, timeout=20)

            datos = resp.json()
            if datos["status"] == "success" and datos["data"]["result"]:
                for serie in datos["data"]["result"]:
                    df = pd.DataFrame(serie["values"], columns=["tiempo", "valor"])
                    df["tiempo"] = pd.to_datetime(df["tiempo"], unit="s")
                    df["valor"] = pd.to_numeric(df["valor"])
                    if "volume" in serie["metric"]:
                        df["serie"] = f'Disco Usado (%) ({serie["metric"]["volume"]})'
                    else:
                        df["serie"] = ETIQUETAS_METRICAS.get(metrica, metrica)
                    dfs.append(df)

        if dfs:
            combinado = pd.concat(dfs)

            # Grafico en pantalla
            if tipo_grafico == "Línea":
                fig = px.line(combinado, x="tiempo", y="valor", color="serie",
                              title=f"Métricas para {instancia} (Gráfico {idx+1})",
                              labels={"valor": "% uso / KB/s", "tiempo": "Tiempo"})
            else:
                fig = px.bar(combinado, x="tiempo", y="valor", color="serie",
                             title=f"Métricas para {instancia} (Gráfico {idx+1})",
                             labels={"valor": "% uso / KB/s", "tiempo": "Tiempo"})

            fig.update_xaxes(tickformat="%d-%m %H:%M", tickangle=-45, title_text="Tiempo")
            fig.update_yaxes(title_text="% uso / KB/s")

            if any("(%)" in s for s in combinado["serie"].unique()):
                fig.update_yaxes(range=[0, 100])

            st.plotly_chart(fig, use_container_width=True)

            # grafico exportable con tiempo formateado fechas 
            export_df = combinado.copy()
            export_df["tiempo"] = export_df["tiempo"].dt.strftime("%d-%m %H:%M")

            if tipo_grafico == "Línea":
                fig_export = px.line(export_df, x="tiempo", y="valor", color="serie",
                                     title=f"Métricas para {instancia} (Gráfico {idx+1})",
                                     labels={"valor": "% uso / KB/s", "tiempo": "Tiempo"})
            else:
                fig_export = px.bar(export_df, x="tiempo", y="valor", color="serie",
                                    title=f"Métricas para {instancia} (Gráfico {idx+1})",
                                    labels={"valor": "% uso / KB/s", "tiempo": "Tiempo"})

            fig_export.update_layout(width=800, height=500, margin=dict(l=50, r=50, t=50, b=100))

            buf = io.BytesIO()
            fig_export.write_image(buf, format="png", scale=2)
            img_bytes = buf.getvalue()
            imagenes_graficos.append((f"Gráfico {idx+1}", img_bytes))
        else:
            st.warning(f"No se encontraron datos para el gráfico {idx+1}.")


    # EXPORTACIÓN PDF y WORD

    if imagenes_graficos:
        fecha_actual = datetime.now().strftime("%d/%m/%Y %H:%M")

        # PDF
        pdf_buf = io.BytesIO()
        c = canvas.Canvas(pdf_buf, pagesize=letter)
        width, height = letter

        c.setFont("Helvetica-Bold", 20)
        c.drawCentredString(width/2, height-100, "Izytech - Reporte de Métricas")
        c.setFont("Helvetica", 12)
        c.drawCentredString(width/2, height-130, f"Generado el {fecha_actual}")
        try:
            logo = "logo_izytech.png"
            c.drawImage(logo, width/2-60, height-250, width=120, height=120, preserveAspectRatio=True)
        except:
            pass
        c.showPage()

        for titulo, img_bytes in imagenes_graficos:
            img = Image.open(io.BytesIO(img_bytes))
            max_width = width - 100
            max_height = height - 200
            aspect_ratio = img.height / img.width

            new_width = max_width
            new_height = int(new_width * aspect_ratio)
            if new_height > max_height:
                new_height = max_height
                new_width = int(new_height / aspect_ratio)

            x_pos = (width - new_width) / 2
            y_pos = (height - new_height) / 2

            img_stream = io.BytesIO()
            img.save(img_stream, format="PNG")
            img_stream.seek(0)

            c.setFont("Helvetica-Bold", 14)
            c.drawCentredString(width/2, height - 50, titulo)

            c.drawImage(ImageReader(img_stream),
                        x_pos, y_pos,
                        width=new_width, height=new_height,
                        preserveAspectRatio=True)

            c.showPage()

        c.save()
        st.download_button("⬇️ Descargar PDF", data=pdf_buf.getvalue(),
                           file_name="reporte_izytech.pdf", mime="application/pdf")

        # Word
        doc = Document()
        doc.add_heading("Izytech - Reporte de Métricas", 0)
        doc.add_paragraph(f"Generado el {fecha_actual}")
        doc.add_paragraph(" ")

        for titulo, img_bytes in imagenes_graficos:
            doc.add_heading(titulo, level=1)
            img_stream = io.BytesIO(img_bytes)
            doc.add_picture(img_stream, width=Inches(6))

        word_buf = io.BytesIO()
        doc.save(word_buf)
        st.download_button("⬇️ Descargar Word", data=word_buf.getvalue(),
                           file_name="reporte_izytech.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
